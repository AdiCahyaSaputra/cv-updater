#!/usr/bin/env python3

import argparse
import os
from docx import Document
from PyPDF2 import PdfReader, PdfWriter
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

class CVModifier:
    def __init__(self, input_file):
        self.input_file = input_file
        self.file_type = os.path.splitext(input_file)[1].lower()

    def modify_docx(self, replacements):
        f_docx = open(self.input_file, 'rb')
        doc = Document(f_docx)
        f_docx.close()

        output_file = f'modified_{os.path.basename(self.input_file)}'

        # Replace text in paragraphs while preserving formatting
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                placeholder = '${' + key + '}'
                if placeholder in paragraph.text:
                    print(f"Found placeholder: {key}")
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))

        # Replace text in tables while preserving formatting
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            placeholder = '${' + key + '}'
                            if placeholder in paragraph.text:
                                for run in paragraph.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, str(value))

        doc.save(output_file)
        return output_file

    def modify_pdf(self, replacements):
        reader = PdfReader(self.input_file)
        writer = PdfWriter()
        output_file = f'modified_{os.path.basename(self.input_file)}'

        for page_num in range(len(reader.pages)):
            # Create a new PDF with replacements
            packet = BytesIO()
            can = canvas.Canvas(packet, pagesize=letter)
            page = reader.pages[page_num]
            
            # Extract text and find placeholder positions
            # Note: This is a simplified approach. For more accurate positioning,
            # you might need to use more sophisticated PDF text extraction
            text_content = page.extract_text()
            
            for key, value in replacements.items():
                placeholder = f'{{${key}}}'
                if placeholder in text_content:
                    # For demonstration, we're placing the new text at a fixed position
                    # In a real implementation, you'd want to calculate the exact position
                    can.drawString(100, 100, f"{key}: {value}")
            
            can.save()
            packet.seek(0)
            new_pdf = PdfReader(packet)
            
            # Merge original page with new content
            page.merge_page(new_pdf.pages[0])
            writer.add_page(page)

        with open(output_file, 'wb') as output_file_handle:
            writer.write(output_file_handle)

        return output_file

    def modify(self, replacements):
        if self.file_type == '.docx':
            return self.modify_docx(replacements)
        elif self.file_type == '.pdf':
            return self.modify_pdf(replacements)
        else:
            raise ValueError(f"Unsupported file type: {self.file_type}")

def main():
    parser = argparse.ArgumentParser(
        description='A tool for dynamically modifying CV content in DOCX and PDF files.',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
            Examples:
            # Update years of experience
            python cv_modifier.py template.docx --experience "5 years"

            # Update multiple skills
            python cv_modifier.py template.docx --skills "Python, JavaScript, Docker"

            # Update current position
            python cv_modifier.py template.pdf --position "Senior Software Engineer"

            # Use custom placeholders
            python cv_modifier.py template.docx --custom location "New York" --custom salary "$120,000"

            # Combine multiple updates
            python cv_modifier.py template.docx --experience "7 years" --skills "Python, AWS" --position "Tech Lead"

            Placeholder Format:
            In your CV template, use ${placeholder} format, e.g., ${experience}, ${skills}, ${position}
            For custom placeholders, the format is ${key} where key is your custom key name
        ''')

    parser.add_argument('input_file',
                        help='Path to the CV template file (.docx or .pdf)')
    parser.add_argument('--experience',
                        help='Specify your years of experience (e.g., "5 years", "3+ years")')
    parser.add_argument('--skills',
                        help='List your skills, separated by commas (e.g., "Python, JavaScript, Docker")')
    parser.add_argument('--custom', nargs=2, action='append',
                        metavar=('KEY', 'VALUE'),
                        help='Add custom placeholder replacements. Can be used multiple times. Format: --custom KEY VALUE')


    args = parser.parse_args()

    # Validate input file
    if not os.path.exists(args.input_file):
        print(f"Error: File '{args.input_file}' not found")
        return

    # Build replacements dictionary
    replacements = {
        "skills": "PHP & Laravel, Docker, Linux, Git, PostgreSQL, Typescript, NextJS, React, NuxtJS, Vue, SvelteKIt, React Native, Flutter, Python",
    }
    if args.experience:
        replacements['experience'] = args.experience
    if args.skills:
        replacements['skills'] = replacements['skills'] + ", " + args.skills
    if args.custom:
        for key, value in args.custom:
            replacements[key] = value

    print(replacements)

    try:
        modifier = CVModifier(args.input_file)
        output_file = modifier.modify(replacements)
        print(f"Modified CV saved as: {output_file}")
    except Exception as e:
        print(f"Error: {str(e)}")

if __name__ == '__main__':
    main()